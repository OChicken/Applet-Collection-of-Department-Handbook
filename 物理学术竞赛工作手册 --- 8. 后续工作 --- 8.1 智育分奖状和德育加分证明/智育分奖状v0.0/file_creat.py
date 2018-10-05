#coding=utf-8
import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
#Main Body
def get_time(year,month):
    return


def get_name(name_string, record):
    check = True
    length = len(name_string)

    while check == True and length > 0:
        point = name_string.find('、')
        length = len(name_string)

        if point >0 :
            record.append(name_string[0:point])
            name_string = name_string.replace(name_string[0:point+1],'')
        elif length > 0:
            record.append(name_string[0:length])
            name_string = name_string.replace(name_string[0:length+1],'')
        else:
            check = False

    return

def exchange_result(record):
    output = ''

    temp = record[0]
    del record[0]
    record.append(temp)


    for num in range(0,len(record)):
        output = output + record[num] + '、'
    return output[0:len(output)-1]

def file_creat(name,fill,price):
    #Exchange name
    paragraph = file.add_paragraph('\n')
    run = paragraph.add_run(fill + '   同学：')
    run.font.size = Pt(22)
    run.font.name = u'华文新魏'
    run.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')

    paragraph = file.add_paragraph('\n')
    run = paragraph.add_run('荣获华南理工大学第五届大学生物理学术竞赛')
    run.font.size = Pt(22)
    run.font.name = u'仿宋'
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

    paragraph = file.add_paragraph('')
    run = paragraph.add_run('SCUT Undergraduate Physicists’ Tournament')
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = 'Times New Roman'
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = file.add_paragraph('')
    run = paragraph.add_run(price)
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = u'华文新魏'
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')

    paragraph = file.add_paragraph('')
    run = paragraph.add_run('教务处   校团委   物理与光电学院\n二〇一八年十二月')
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = u'宋体'
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    file.save(data_dir +'/' + name +'.docx')
    file._body.clear_content()


name_record = []
name_string = ''
data_dir = os.getcwd()
file = Document()

#Data Input
name_string = str(input('Tell me their names:'))
price = str(input('Price:'))
get_name(name_string, name_record)
size = len(name_record)
#Data Output
for times in range(0, size):
    fill_string = ''
    fill_string = exchange_result(name_record)
    file_creat(name_record[0], fill_string, price)

print('Creat Successfully!')
