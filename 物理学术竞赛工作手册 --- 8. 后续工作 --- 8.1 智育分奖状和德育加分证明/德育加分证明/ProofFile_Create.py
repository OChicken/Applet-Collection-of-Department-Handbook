#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Date    : 2018-09-27 21:27:16
# @Author  : Ma Seoyin (Ma.Seoyin@gmail.com)
# @Link    : https://github.com/OChicken
# @Version : V2

import os
import time
import docx
import xlrd
import win32com.client
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import comtypes.client
from shutil import copyfile


def delEmptyElement(List):
    while List[-1] == '':
        del List[-1]
    return List


def WriteDocxFiles(template, NameDisplay, k):
    template.paragraphs[3 + 20 * k].clear()
    run = template.paragraphs[3 + 20 * k].add_run(NameDisplay)
    run.font.size = Pt(12)
    run.font.name = u'微软雅黑'
    run.bold = False
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')


# 欢迎辞
Dir = os.getcwd() + '/'
word = comtypes.client.CreateObject('Word.Application')
doc_app=win32com.client.Dispatch('Word.Application')
print('欢迎使用 物理学术竞赛德育加分证明自动生成 小程序 :)\n'
      '<<<<<<<<<<<<<<<<<<<<< 啦啦啦我是分割线 >>>>>>>>>>>>>>>>>>>>>\n'
      'author:\n马守然 (2014级应用物理学)\n学术科创部\n物理与光电学院团委学生会\nEmail: 1941688873@qq.com / Ma.Seoyin@gmail.com\n'
      'Link: https://github.com/OChicken\n'
      '<<<<<<<<<<<<<<<<<<<<< 啦啦啦我是分割线 >>>>>>>>>>>>>>>>>>>>>\n'
      '德育分0.5加分证明生成中, 请稍候......')
ProofList = xlrd.open_workbook(Dir + 'Proof.xlsx')
Start = time.clock()
# 新建德育加分0.5分的文件夹
ProofDir = Dir + '德育分0.5/'
if os.path.exists(ProofDir.rstrip('/')) == False:
    os.makedirs(ProofDir)
ProofTemplate = ProofDir + 'template_0.5.docx'
copyfile('template_0.5.docx', ProofTemplate)
Proof_05 = ProofList.sheet_by_index(0)
Name = Proof_05.col_values(0)  # 和名字序列
size = len(Name)
for k in range(size - 1):
    if k == 0:
        doc=doc_app.Documents.Open(Dir + 'template_0.5.docx')
        doc.Content.Copy()
        sel = doc_app.Selection
        sel.Range.Paste()
        doc.SaveAs(ProofTemplate)
        doc.Close()
    else:
        doc=doc_app.Documents.Open(ProofTemplate)
        sel = doc_app.Selection
        sel.Range.Paste()
        doc.SaveAs(ProofTemplate)
        doc.Close()
template = docx.Document(ProofTemplate)
FileName = ProofDir + '德育分0.5加分证明'
for k in range(size):
    WriteDocxFiles(template, Name[k], k)
template.save(FileName + '.docx')
doc = word.Documents.Open(FileName + '.docx')
doc.SaveAs(FileName + '.pdf', FileFormat=17)
doc.Close()
os.remove(ProofTemplate)
End = time.clock()
print('所有德育分0.5加分证明已生成, 用时' + str(End-Start) + '秒')
print('德育分1.0加分证明生成中, 请稍候......')
Start = time.clock()
# 新建德育加分1.0分的文件夹
ProofDir = Dir + '德育分1.0/'
if os.path.exists(ProofDir.rstrip('/')) == False:
    os.makedirs(ProofDir)
ProofTemplate = ProofDir + 'template_1.0.docx'
copyfile('template_1.0.docx', ProofTemplate)
Proof_10 = ProofList.sheet_by_index(1)
for j in range(Proof_10.ncols):
    Team = str(int(Proof_10.col_values(j)[0]))
    Name = Proof_10.col_values(j)[1:]
    delEmptyElement(Name)
    size = len(Name)
    # 生成名字序列那么长的奖状页数 (譬如说一个队有5人, 生成的奖状页数就是5页)
    for k in range(size - 1):
        if k == 0:
            doc=doc_app.Documents.Open(Dir + 'template_1.0.docx')
            doc.Content.Copy()
            sel = doc_app.Selection
            sel.Range.Paste()
            doc.SaveAs(ProofTemplate)
            doc.Close()
        else:
            doc=doc_app.Documents.Open(ProofTemplate)
            sel = doc_app.Selection
            sel.Range.Paste()
            doc.SaveAs(ProofTemplate)
            doc.Close()
    template = docx.Document(ProofTemplate)
    Leader = Name[0]  # 队长名字
    FileName = ProofDir + Team
    for k in range(size):
        WriteDocxFiles(template, Name[k], k)
    template.save(FileName + '.docx')
    doc = word.Documents.Open(FileName + '.docx')
    doc.SaveAs(FileName + '.pdf', FileFormat=17)
    doc.Close()
    os.remove(ProofTemplate)
End = time.clock()
print('所有德育分1.0加分证明已生成, 用时' + str(End-Start) + '秒')
input('请按回车关闭本宝宝 :)')
